import openai
from groq import Groq
import speech_recognition as sr
import pyttsx3
import webbrowser
import subprocess
import os
import sys
import psutil
from psutil._common import bytes2human
import wmi
import time
import socket
import pandas as pd

datosDePc = ""
infoCpu = "" #para guardar la info
infoMemoria = "" #Pa guardar la info
infoServicios = "" #Pa guardar la info 
infoRed  = "" #Pa guardar la info 
infoprocesos = "" #Pa guardar la info 
#APIkEY
openai.api_key = "" #La key personal de Open AI, si es una app ya lanzada aqui iria el api de la empresa

#Inicializacion Voz
engine = pyttsx3.init()

#PASAR AUIDO A TEXTO
def transcribe_audio_to_text(filename): #Pasa audio a texto
    recognizer = sr.Recognizer() 
    try:
        with sr.AudioFile(filename) as source: #Reconoce el audio capturado por el micro usando este audio como fuente
            audio = recognizer.record(source)
        return recognizer.recognize_google(audio) #Ya esto es la funcion de google ahi no nos metamos mucho a explicarlo LOL
    except sr.UnknownValueError:
        print("No se pudo entender el audio.")
    except sr.RequestError as e:
        print(f"Error al solicitar resultados desde el servicio de Google Speech Recognition: {e}")
    except Exception as e:
        print(f"Error inesperado: {e}")
    return None

#Generacion y comportamiento de la IA
def generate_response(prompt):
    global datosDePc
    global infoCpu
    global infoMemoria
    global infoServicios
    global infoRed
    global infoprocesos
    try:
        if "abrir pagina" in prompt:
            url = prompt.replace("abrir pagina ", "").strip()
            webbrowser.open_new_tab(url)
            return f"Abriendo {url}"
        
        elif "info" in prompt:
            #--------------------------------VAR NECESARIAS----------------------------
            numeroproceso = 0 #Llevar orden de los procesos
            respuesta = "" #Respuesta
            #--------------------------------CPU---------------------------------------
            infoCpu = "" #para guardar la info
            c = wmi.WMI() #wmi
            cpu = c.Win32_Processor()[0] #Procesador 0, esto es pensando sea una build de solo un cpu, pero se podria modificar esta linea para un server 
            miCpu = cpu.Name #Nombre de cpu
            miCpuArquitectura = cpu.Architecture #Arquitectura en nanometros
            miUsoCpu = psutil.cpu_percent(interval=1) #Porcentaje de uso
            miCpuThreads = psutil.cpu_count() #Hilos
            miCpuCores = psutil.cpu_count(logical=False) #Nucleos
            
            infoCpu = ( #Muestra de informacion
                     f"Detalles de cpu:\n"
                     f"Nombre de CPU: {miCpu}\n"
                     f"Arquitectura del CPU: {miCpuArquitectura}nm\n"
                     f"Uso de la CPU: {miUsoCpu}%\n"
                     f"Núcleos físicos: {miCpuCores}\n"
                     f"Hilos totales: {miCpuThreads}\n\n"
                 )
            #--------------------------------MEMORIA---------------------------------------
            infoMemoria = "" #Pa guardar la info
            miMemoria = psutil.virtual_memory() #Acceso a memoria, cantidad de memoria
            miMemoria = bytes2human(miMemoria[0]) #Pasar a formato legible
            miMemoriaPorctEnUso = psutil.virtual_memory()[2] #Porcentaj de uso
            miMemoriaEnUso = psutil.virtual_memory()[3]/1000000000 #Gigas en uso
            miMemoriaPorctEnUso = str(miMemoriaPorctEnUso) + "%"
            miMemoriaEnUso = str(miMemoriaEnUso) + "G"
            infoMemoria = ( #Muestra de informacion
                     f"Detalles de memoria:\n"
                     f"Memoria total: {miMemoria}\n"
                     f"Porcentaje de memoria en uso: {miMemoriaPorctEnUso}\n"
                     f"Cantidad de gigas en uso: {miMemoriaEnUso}\n\n"
                 )
            #--------------------------------SERVICIOS DE WINDOWS---------------------------------------
            #pos la ruta
            excel_path = "DiagnosticoDePcPorAI//Data//servicestodisable.xlsx"
            
            # #df es el excel
            df = pd.read_excel(excel_path)
            
            #Normalizar exce
            df.columns = df.columns.str.strip()
            df['Service Name'] = df['Service Name'].str.strip()
            
            numeroServicio = 0
            misServiciosWin = psutil.win_service_iter()
            
            print("Buscando servicios en ejecución anomalos que se pueden deshabilitar:\n")
            time.sleep(0.5) #Se da un tiempo para la correcta carga del excl y psuntil   
            
            for servicio in misServiciosWin:
                if servicio.status() == "running": #Si el servicio esta ejecutandose
                    name = servicio.name()  #Nombre servicio
                    pid = servicio.pid()  #PID servicio
                    start_type = servicio.start_type()  #Tipo de inicio
                    description = servicio.display_name()  #Desc
                    
                    match = df[df['Service Name'] == name] #Se usara como match el excel con los servicios que se pueden deshabilitar
                    if not match.empty and match.iloc[0]['Recommendation'] == "OK to disable": #Si hay match y se puede deshabilitar
                        numeroServicio += 1 #Se suma uno a los servicios
                        
                        infoServicios = ( #Se guarda la info del servicio
                                         f"Servicio número {numeroServicio}: {name}\n"
                                         f"PID: {pid if pid is not None else 'Sin PID'}\n"
                                         f"Estado del servicio: activo\n"
                                         f"Tipo de inicio: {start_type}\n"
                                         f"Descripción del servicio: {description}\n"
                                         f"Recomendación: {match.iloc[0]['Recommendation']}\n\n"
                                         )
            if infoServicios == "": ##Si no hay info
                infoServicios = "No se encuentran servicios anomalos" #Si no hay pues se dice que no hay
            #--------------------------------RED---------------------------------------
            infoRed  = ""
            miRed = psutil.net_io_counters(pernic=True) #Detalles de trafico en red
            misInterfacesRed = psutil.net_if_addrs() #Detalles de interfaces de red, detalles de direciones Ip Y Mac
            misEstadisticasRed = psutil.net_if_stats() #Detalles de estadisticas de red, actividad de interfaz y vel
            
            for interface, stats in miRed.items():
                bytesEnviados = stats.bytes_sent #Calcula los bytes que entran por puerto de red
                bytesRecibidos = stats.bytes_recv #Calcula los bytes que salen por puerto de red
                paquetesEnviados = stats.packets_sent #Calcula los paquetes que entran por puerto de red
                paquetesRecibidos  = stats.packets_recv #Calcula los paquetes que salen por puerto de red
                
                if interface in misInterfacesRed:
                    direcionesIp = misInterfacesRed[interface] #Lista de direcciones ip
                    for direccion in direcionesIp:
                        if direccion.family == socket.AF_INET: #pvv4
                            ip = direccion.address
                            mascara = direccion.netmask
                        elif direccion.family == socket.AF_INET6: #ipv6
                            ipv6 = direccion.address
                            mascara_ipv6 = direccion.netmask
                        elif direccion.family == psutil.AF_LINK: #MAC address
                            mac = direccion.address
                            
                if interface in misEstadisticasRed:
                    estado = misEstadisticasRed[interface]
                    upOrDown = "Up" if estado.isup else "Down" #Estado de puerto de red
                    velocidad = f"{estado.speed} Mbps" if estado.speed else "Desconocida" #Velocidad del puerto de red
                    
                infoRed = ( #Muestra de informacion
                     f"Detalles de la red - {interface}:\n"
                     f"IP: {ip if 'ip' in locals() else 'No disponible'}\n"
                     f"Máscara: {mascara if 'mascara' in locals() else 'No disponible'}\n"
                     f"IPv6: {ipv6 if 'ipv6' in locals() else 'No disponible'}\n"
                     f"MAC: {mac if 'mac' in locals() else 'No disponible'}\n"
                     f"Estado activo: {upOrDown}\n"
                     f"Velocidad: {velocidad}\n"
                     f"Bytes enviados: {bytesEnviados} bytes\n"
                     f"Bytes recibidos: {bytesRecibidos} bytes\n"
                     f"Paquetes enviados: {paquetesEnviados}\n"
                     f"Paquetes recibidos: {paquetesRecibidos}\n\n"
                 )
                infoRedApi = ( #Muestra de informacion
                     f"Detalles de la red - {interface}:\n"
                     f"Estado activo: {upOrDown}\n"
                     f"Velocidad: {velocidad}\n"
                     f"Bytes enviados: {bytesEnviados} bytes\n"
                     f"Bytes recibidos: {bytesRecibidos} bytes\n"
                     f"Paquetes enviados: {paquetesEnviados}\n"
                     f"Paquetes recibidos: {paquetesRecibidos}\n\n"
                 )
            #--------------------------------PROCESOS EN CURSO---------------------------------------
            for proceso in psutil.process_iter(): #Se llama la verificacion del gpu por primera vez
                proceso.cpu_percent() #Esto es por el como psutil calcula este uso de cpu
            print("Verificando los procesos en curso")
            time.sleep(1) #Se da un tiempo para que psutil se asegure del %           
            
            procesosSistema = psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'io_counters']) #Obtener la info de los procesos
            topProcesos = sorted(procesosSistema, key=lambda p: p.info['cpu_percent'], reverse=True)[:10] #Ordenar los 10 primeros procesos por uso de CPU
            for process in topProcesos: #For de los 10 procesos
                numeroproceso += 1 #Se incrementa el contador
                name = process.info['name'] #Nombre de proceso
                porcentajeCpu = process.info['cpu_percent'] #Procentaje en CPU
                procentajeMemoria = process.info['memory_percent'] #Porcentaje en memoria
                verificarDisco = process.info.get('io_counters', None) #Uso de disco
                if verificarDisco: #Si se esta usando el disco
                    usodeldisco = verificarDisco.read_bytes + verificarDisco.write_bytes #info de uso
                else:
                    usodeldisco = "Sin detalle de uso" #Sin uso de disco xdxd
                
                procesoEnCurso = (
                    f"Proceso número {numeroproceso}: {name}\n"
                    f"Uso de CPU: {porcentajeCpu}%\n"
                    f"Uso de Memoria: {procentajeMemoria}%\n"
                    f"Uso de Disco: {usodeldisco} bytes\n\n"
                )
                respuesta += procesoEnCurso
                infoprocesos = respuesta
                response = respuesta
                
            #--------------------------------IMPRESION Y RESPUESTA---------------------------------------
            response = "Datos del Cpu: \n" + infoCpu +  "Datos de memoria: \n" + infoMemoria + "Datos de red: \n" + infoRed + "Datos de servicios: \n" + infoServicios + "Datos de procesos: \n" + response
            datosDePc = "Datos del Cpu: \n" + infoCpu +  "Datos de memoria: \n" + infoMemoria + "Datos de red: \n" + infoRedApi + "Datos de servicios: \n" + infoServicios + "Datos de procesos: \n" + response
            return response
        elif "salir" in prompt:
            mensajeSalida = "Gracias por utilizar nuestro sistema."
            print(mensajeSalida)
            sys.exit(0)
            return mensajeSalida
        else: #Aqui se tiene que enviar la info de sistema a GroqAi, deberia ser asi
            client = Groq(api_key="")
            mensaje = prompt + "\n Los de detalles de mi PC son: \n\n" + datosDePc
            chat_completion = client.chat.completions.create(messages=[
                {
                    "role": "system",
                    "content": "Eres un sistema de diagnostico de computadores, Brindaras apoyo a un p´roblema que presente el usuario en base a la informacion de la computadora del mismo. Esto lo haras de manera precisa verificando cada detalle que te brinde el cliente de su pc. En la respuesta daras un analisis general de la PC asi como recomendaciones."
                    },
                {
                    "role": "user",
                    "content": mensaje
                    }
                ],
                                                             model="llama3-8b-8192",
                                                             temperature=0.5
                                                             )
            resultado = chat_completion.choices[0].message.content
            if resultado:
                from docx import Document
                from docx.shared import Cm
                import os
                
                document = Document()
                style = document.styles['Normal']
                style.font.name = 'Calibri'
                header_section = document.sections[0]
                header = header_section.header
                header_text = header.paragraphs[0]
                header_text.text = "--DIAGNOSTICOS PC: Grupo 1--"
                
                document.add_heading('Analisis de Pc mediante AI', 0)
                
                document.add_picture('DiagnosticoDePcPorAI\\ai.png', width=Cm(5))
                
                p = document.add_paragraph('Se muestra la informacion extraida de la PC:')
                
                document.add_heading('CPU')
                p = document.add_paragraph(infoCpu)
                
                document.add_heading('Memoria')
                p = document.add_paragraph(infoMemoria)
                
                document.add_heading('Servicios')
                p = document.add_paragraph(infoServicios)
                
                document.add_heading('Red')
                p = document.add_paragraph(infoRed)
                
                document.add_heading('Procesos')
                p = document.add_paragraph(infoprocesos)
                
                document.add_page_break()
                
                document.add_heading('Recomendaciones y Analisis')
                p = document.add_paragraph(resultado)
                
                
                document.save('DiagnosticoDePcPorAI\\Analisis.docx')
                
                from docx2pdf import convert
                docx_file = 'DiagnosticoDePcPorAI\\Analisis.docx'
                pdf_file = 'DiagnosticoDePcPorAI\\Analisis.pdf'
                
                convert(docx_file, pdf_file)
            else:
                print("No se logro optener respuesta de la AI")
        return(chat_completion.choices[0].message.content)
    except Exception as e:
        return f"Error al generar la respuesta: {e}"

#Texto de la IA a Voz
def speak_text(text):
    try:
        engine.say(text)
        engine.runAndWait()
    except Exception as e:
        print(f"Error al convertir texto a voz: {e}")

#Se verifica el user diga "Sistema" para empezar a verificar
def listen_for_keyword():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        print("Diga 'Sistema' para comenzar a verificar el sistema...")
        audio = recognizer.listen(source)
        try:
            transcription = recognizer.recognize_google(audio)
            return transcription.lower() == "sistema"
        except sr.UnknownValueError:
            print("No se entendió la palabra clave.")
        except sr.RequestError as e:
            print(f"Error con el servicio de reconocimiento de voz: {e}")
        except Exception as e:
            print(f"Error inesperado: {e}")
    return False

#Obtencion del audio
def record_audio(filename):
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.pause_threshold = 1  #Sleep
        print("En espera de orden...")
        audio = recognizer.listen(source)
        try:
            with open(filename, "wb") as f:
                f.write(audio.get_wav_data())
            return True
        except Exception as e:
            print(f"Error al grabar el audio: {e}")
            return False

#Funcion principal
def main():
    while True:
        if listen_for_keyword():
            filename = "input.wav"
            if record_audio(filename):
                text = transcribe_audio_to_text(filename) #Transcribe el audio a texto
                if text:
                    print(f"Dijiste: {text}")
                    response = generate_response(text) #Muestra el audio que optiene del micro pasado a texto
                    print(f"Ai dice: {response}") #Muestra la respuesta de la AI
                    speak_text(response)

if __name__ == "__main__":
    main()
